"""
Parser untuk file DOCX RPS yang diupload ulang oleh dosen (hasil isian dari
file yang sebelumnya di-download dari sistem).

Dua mode:
  - deterministik (extract_detail_table): baca tabel "Detail Perkuliahan 16
    Minggu" berdasarkan POSISI kolom, bukan nama header -- tabel ini punya
    header 2-baris dengan merged cell yang tidak bisa dipetakan lewat nama
    kolom. Cepat, tanpa AI, tapi hanya mencakup 16 minggu.
  - AI (extract_full_json_via_ai): kirim seluruh teks dokumen ke LLM, minta
    JSON lengkap sesuai skema rps_bp (lihat
    bp_edu_document/utils/json_importer.py) yang lalu diproses oleh
    import_rps() yang sama dipakai wizard Import JSON.
"""
import base64
import io
import json
import logging
import re

from docx import Document

_logger = logging.getLogger(__name__)

_DETAIL_COLS = ['minggu', 'deskripsi', 'indikator', 'kriteria', 'tatap_muka', 'daring', 'materi', 'bobot']

_MAX_DOC_TEXT_CHARS = 24000


# ─── Validasi file ───────────────────────────────────────────────────────────

def load_docx(docx_file_b64, filename):
    """Decode + buka file .docx. Return (Document|None, error_message|None)."""
    if not filename or not filename.lower().endswith('.docx'):
        return None, f'Hanya file .docx yang didukung (file yang diupload: {filename or "(tanpa nama)"}).'
    try:
        raw = base64.b64decode(docx_file_b64)
    except Exception:
        return None, 'File tidak bisa dibaca (isi base64 tidak valid).'
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as e:
        return None, f'File bukan dokumen Word (.docx) yang valid, atau rusak: {e}'
    return doc, None


# ─── Mode deterministik: tabel Detail 16 Minggu ─────────────────────────────

def _find_detail_table(doc):
    """Cari tabel Detail Perkuliahan 16 Minggu lewat penanda 'Sub-CPMK' di baris headernya."""
    for table in doc.tables:
        for row in table.rows[:4]:
            row_text = ' '.join(c.text for c in row.cells)
            if 'sub-cpmk' in row_text.lower():
                return table
    return None


def extract_detail_table(doc):
    """
    Ekstrak tabel Detail Perkuliahan 16 Minggu berdasarkan posisi kolom.

    Return dict: {'ok': bool, 'weeks': {1: {...}, ...}, 'errors': [...], 'warnings': [...]}
    'ok' True hanya bila seluruh 16 minggu terbaca bersih (tanpa error/warning).
    """
    table = _find_detail_table(doc)
    if table is None:
        return {
            'ok': False,
            'weeks': {},
            'errors': [
                'Tabel "Detail Perkuliahan 16 Minggu" (kolom Sub-CPMK) tidak ditemukan. '
                'Pastikan ini file RPS hasil download dari sistem, dan struktur tabelnya '
                'tidak diubah (baris/kolom tidak dihapus atau digabung).',
            ],
            'warnings': [],
        }

    weeks = {}
    warnings = []
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        first = cells[0].text.strip()
        # Kolom "Mg Ke-" bisa berisi angka polos ("1") atau format "1/16"
        # (dipakai template saat ini).
        match = re.match(r'^(\d+)\s*(?:/\s*16)?$', first)
        if not match:
            continue
        minggu = int(match.group(1))
        if not (1 <= minggu <= 16):
            continue
        if len(cells) < 8:
            warnings.append(
                f'Minggu {minggu}: baris tabel hanya punya {len(cells)} kolom '
                f'(diharapkan 8), dilewati.'
            )
            continue
        weeks[minggu] = {
            _DETAIL_COLS[i]: (minggu if i == 0 else cells[i].text.strip())
            for i in range(8)
        }

    errors = []
    missing = [w for w in range(1, 17) if w not in weeks]
    if missing:
        errors.append('Minggu yang tidak terbaca: ' + ', '.join(str(w) for w in missing) + '.')

    return {
        'ok': not errors and not warnings,
        'weeks': weeks,
        'errors': errors,
        'warnings': warnings,
    }


# ─── Mode AI: ekstraksi JSON penuh ───────────────────────────────────────────

RPS_JSON_SCHEMA_PROMPT = """Kamu adalah ekstraktor data akademik. Tugasmu: baca teks dokumen RPS
(Rencana Pembelajaran Semester) di bawah, lalu keluarkan HANYA JSON valid
(tanpa markdown code fence, tanpa penjelasan tambahan) sesuai skema berikut.
Field yang tidak ditemukan di dokumen, isi dengan string kosong "" atau list
kosong []. JANGAN mengarang data yang tidak benar-benar ada di dokumen.

Skema JSON:
{
  "meta": {
    "kode_mk": string, "nama_mk": string,
    "sks_teori": integer, "sks_praktik": integer, "semester": integer,
    "status": "Wajib" atau "Pilihan", "kategori": string,
    "tanggal_penyusunan": "YYYY-MM-DD" atau "",
    "dosen_pengampu": string
  },
  "deskripsi_singkat": string,
  "bahan_kajian": string,
  "matakuliah_syarat": string,
  "cpl_prodi": [{"kode": string, "tipe": string, "deskripsi": string}],
  "cpmk": [{"kode": "CPMK-1", "cpl": "kode CPL dipisah koma", "deskripsi": string}],
  "sub_cpmk": [{"kode": string, "cpmk": "CPMK-1", "deskripsi": string, "minggu": string, "taksonomi": string}],
  "korelasi": [{"cpmk": "CPMK-1", "m1": "x atau \\"\\"", "m2": "...", "...": "...", "m12": "..."}],
  "korelasi_cpl": [{"sub": string, "p1": "x atau \\"\\"", "...": "...", "p20": "...", "bobot": string, "minggu": string}],
  "penilaian": [{"jenis": string, "bobot": string, "c1": "x atau \\"\\"", "...": "...", "c10": "..."}],
  "detail": [{"minggu": integer, "deskripsi": string, "indikator": string, "kriteria": string, "tatap_muka": string, "daring": string, "materi": string, "bobot": string}],
  "pustaka_utama": [{"kode": string, "referensi": string}],
  "pustaka_pendukung": [{"kode": string, "referensi": string}]
}

"detail" idealnya berisi 16 item (minggu 1 sampai 16) sesuai isi tabel
"Detail Perkuliahan 16 Minggu" di dokumen. Kolom m1..m12/p1..p20/c1..c10 diisi
"x" hanya bila memang ditandai di tabel korelasi/penilaian, selain itu "".
"""


def dump_doc_text(doc):
    """Ratakan seluruh teks dokumen (paragraf + tabel) jadi satu string untuk dikirim ke AI."""
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        lines.append('--- TABEL ---')
        for row in table.rows:
            lines.append(' | '.join(c.text.strip().replace('\n', ' ') for c in row.cells))
    text = '\n'.join(lines)
    if len(text) > _MAX_DOC_TEXT_CHARS:
        text = text[:_MAX_DOC_TEXT_CHARS] + '\n...[dipotong, dokumen terlalu panjang]'
    return text


def _extract_json_from_text(content):
    content = (content or '').strip()
    m = re.search(r'```(?:json)?\s*(\{.*\})\s*```', content, re.DOTALL)
    if m:
        content = m.group(1)
    else:
        start, end = content.find('{'), content.rfind('}')
        if start != -1 and end != -1 and end > start:
            content = content[start:end + 1]
    try:
        return json.loads(content)
    except Exception:
        return None


def extract_full_json_via_ai(env, doc):
    """
    Kirim teks dokumen ke AI agent aktif, minta JSON sesuai skema rps_bp.
    Return (data_dict|None, error_message|None).
    """
    agent = env['ai.agent'].sudo().search([('active', '=', True)], limit=1)
    if not agent:
        return None, (
            'AI belum dikonfigurasi. Buat AI Agent aktif di '
            'Settings → Technical → AI Agents, atau isi API Key di Settings → General Settings.'
        )

    text = dump_doc_text(doc)
    messages = [
        {'role': 'system', 'content': RPS_JSON_SCHEMA_PROMPT},
        {'role': 'user', 'content': f'--- Teks Dokumen RPS ---\n{text}'},
    ]
    response, err = agent._call_api(messages, use_tools=False)
    if response is None:
        return None, err or 'AI tidak merespons.'
    choices = response.get('choices', [])
    if not choices:
        return None, 'Respons AI tidak berisi choices.'
    content = choices[0].get('message', {}).get('content', '') or ''
    data = _extract_json_from_text(content)
    if data is None:
        _logger.warning('AI upload RPS: gagal parse JSON dari respons: %s', content[:500])
        return None, 'AI tidak mengembalikan JSON yang valid (lihat log server untuk respons mentah).'
    return data, None


def summarize_ai_json(data):
    """Hitung ringkasan per bagian dari JSON hasil AI. Return (ok: bool, counts: dict, weeks_found: set)."""
    detail = data.get('detail', []) or []
    weeks_found = set()
    for d in detail:
        m = str(d.get('minggu', '')).strip()
        if m.isdigit() and 1 <= int(m) <= 16:
            weeks_found.add(int(m))

    cpl_prodi = data.get('cpl_prodi', []) or []
    cpmk = data.get('cpmk', []) or []

    counts = {
        'Detail Minggu': f'{len(weeks_found)} / 16',
        'CPL Prodi': len(cpl_prodi),
        'CPMK': len(cpmk),
        'Sub-CPMK': len(data.get('sub_cpmk', []) or []),
        'Pustaka Utama': len(data.get('pustaka_utama', []) or []),
        'Pustaka Pendukung': len(data.get('pustaka_pendukung', []) or []),
        'Korelasi CPMK-Minggu': len(data.get('korelasi', []) or []),
        'Korelasi Sub-CPMK-CPL': len(data.get('korelasi_cpl', []) or []),
        'Matriks Penilaian': len(data.get('penilaian', []) or []),
    }
    ok = len(weeks_found) >= 10 and len(cpmk) >= 1 and len(cpl_prodi) >= 1
    return ok, counts, weeks_found
