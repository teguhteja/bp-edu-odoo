"""
DOCX Renderer — port dari rps_bp/edit_rps.py
Mengganti placeholder {field}, {list[0].key}, {nested.field} di template DOCX.

Tidak ada dependensi Odoo — pure python-docx + Pillow.
"""
import copy
import re
import io

from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.run import Run


# ─── Placeholder Parser ──────────────────────────────────────────────────────

# Alias nama field supaya template versi baru tetap bisa membaca context lama.
# Kunci = nama di template, nilai = daftar nama alternatif di context.
FIELD_ALIASES = {
    'pangkat': ['pangkat_golongan'],
    'taksonomi': ['cpl'],
    # Context lama belum punya 'label' ('CPL 1 (S-01)'); pakai kodenya saja.
    'label': ['kode'],
}


def _lookup_key(mapping, key):
    """Ambil key dari dict, dengan fallback ke nama alias (lihat FIELD_ALIASES)."""
    if key in mapping:
        return mapping[key]
    for alias in FIELD_ALIASES.get(key, []):
        if alias in mapping:
            return mapping[alias]
    return None


def _parse_placeholder(placeholder: str) -> list:
    """'{cpl_prodi[0].kode}' → ['cpl_prodi', 0, 'kode']"""
    if placeholder.startswith('{') and placeholder.endswith('}'):
        placeholder = placeholder[1:-1]
    placeholder = placeholder.replace(' ', '')
    parts = re.split(r'\.|\[|\]', placeholder)
    parts = [p for p in parts if p]
    return [int(p) if p.isdigit() else p for p in parts]


def _get_value(data: dict, path_parts: list):
    """Navigate nested dict/list by path. Falls back to data['meta'][key] for flat placeholders."""
    if len(path_parts) == 1 and isinstance(data, dict) and path_parts[0] not in data and 'meta' in data:
        value = _lookup_key(data.get('meta', {}), path_parts[0])
        if value is not None:
            return value

    value = data
    try:
        for part in path_parts:
            if isinstance(value, dict):
                value = _lookup_key(value, part)
            elif isinstance(value, list):
                value = value[part]
            else:
                return None
            if value is None:
                return None
        return value
    except (IndexError, KeyError, TypeError):
        return None


# ─── Image helper ────────────────────────────────────────────────────────────

def _set_image_in_front_of_text(picture):
    """Change image wrapping to 'In Front of Text' via XML manipulation."""
    drawing = picture._inline.getparent()
    inline = drawing.find('.//wp:inline', namespaces=drawing.nsmap)
    if inline is None:
        return
    extent = inline.find('.//wp:extent', namespaces=inline.nsmap)
    effectExtent = inline.find('.//wp:effectExtent', namespaces=inline.nsmap)
    docPr = inline.find('.//wp:docPr', namespaces=inline.nsmap)
    cNvGraphicFramePr = inline.find('.//wp:cNvGraphicFramePr', namespaces=inline.nsmap)
    graphic = inline.find('.//a:graphic', namespaces=inline.nsmap)
    if None in [extent, docPr, graphic]:
        return
    anchor_xml = (
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'distT="0" distB="0" distL="114300" distR="114300" simplePos="0" '
        'relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        '</wp:anchor>'
    )
    anchor = parse_xml(anchor_xml)
    anchor.append(extent)
    if effectExtent is not None:
        anchor.append(effectExtent)
    anchor.append(parse_xml(
        '<wp:wrapNone xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"/>'
    ))
    anchor.append(docPr)
    if cNvGraphicFramePr is not None:
        anchor.append(cNvGraphicFramePr)
    anchor.append(graphic)
    drawing.replace(inline, anchor)


# ─── Text replacement ────────────────────────────────────────────────────────

_PLACEHOLDER_RE = re.compile(r'\{[^{}]+\}')
# Placeholder yang menunjuk elemen list, mis. '{cpl_prodi[3].kode}'. Hanya
# placeholder jenis ini yang boleh memicu penghapusan baris/kolom template.
_INDEXED_RE = re.compile(r'\{[^{}]*\[\s*\d+\s*\][^{}]*\}')


def _replace_text(text: str, data: dict) -> str:
    if not text:
        return text

    def replacer(match):
        ph = match.group(0)
        path = _parse_placeholder(ph)
        value = _get_value(data, path)
        if value is None:
            return ph
        if isinstance(value, list):
            return '\n'.join(str(v) for v in value)
        return str(value)

    return _PLACEHOLDER_RE.sub(replacer, text)


def _write_text_with_markdown_italic(paragraph, text):
    """
    Tulis teks ke run pertama paragraf, dan ubah penanda *miring* gaya markdown
    (lazim pada judul buku di daftar pustaka) menjadi italic Word yang sebenarnya.
    """
    first = paragraph.runs[0]
    segments = re.split(r'\*([^*\n]+)\*', text)
    if len(segments) == 1:
        first.text = text
        return

    # Hasil split berselang-seling: teks biasa, teks miring, teks biasa, ...
    first.text = segments[0]
    anchor = first._r
    for i, segment in enumerate(segments[1:], start=1):
        if not segment:
            continue
        new_r = copy.deepcopy(first._r)
        anchor.addnext(new_r)
        anchor = new_r
        run = Run(new_r, paragraph)
        run.text = segment
        run.italic = (i % 2 == 1)


# ─── Paragraph & table processors ───────────────────────────────────────────

_IMAGE_PLACEHOLDERS = {
    '{dosen_sign}': ('dosen_sign', False),
    '{dosen_sign_small}': ('dosen_sign', True),
    '{mahasiswa_sign}': ('mahasiswa_sign', False),
    '{mahasiswa_sign_small}': ('mahasiswa_sign', True),
}


def _process_paragraph(paragraph, data: dict, images: dict):
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return

    # Image placeholder handling
    for placeholder, (key, is_small) in _IMAGE_PLACEHOLDERS.items():
        if placeholder in full_text and images.get(key):
            img_bytes = images[key]
            width = Inches(0.5) if is_small else Inches(1.0)
            for run in paragraph.runs:
                run.text = ''
            run = paragraph.add_run()
            try:
                from PIL import Image
                with Image.open(io.BytesIO(img_bytes)) as img:
                    if img.mode not in ('RGB', 'RGBA'):
                        img = img.convert('RGBA')
                    buf = io.BytesIO()
                    img.save(buf, format='PNG')
                    buf.seek(0)
                    pic = run.add_picture(buf, width=width)
                    _set_image_in_front_of_text(pic)
            except Exception:
                pass
            return

    new_text = _replace_text(full_text, data)
    if new_text == full_text:
        return
    if paragraph.runs:
        # Kosongkan run lainnya lebih dulu supaya run italic yang ditambahkan
        # di bawah tidak ikut terhapus.
        for run in paragraph.runs[1:]:
            run.text = ''
        _write_text_with_markdown_italic(paragraph, new_text)
    else:
        paragraph.add_run(new_text)


def _row_text(row):
    """Gabungan teks seluruh sel pada satu baris tabel."""
    return '\n'.join(cell.text for cell in row.cells)


def _has_merged_cells(table):
    """True bila tabel memakai gridSpan atau vMerge (tidak aman untuk hapus kolom)."""
    tbl = table._tbl
    return bool(tbl.findall('.//' + qn('w:gridSpan')) or tbl.findall('.//' + qn('w:vMerge')))


def _grid_width(grid):
    """Total lebar tblGrid dalam twips, 0 bila tidak diketahui."""
    if grid is None:
        return 0
    total = 0
    for col in grid.findall(qn('w:gridCol')):
        try:
            total += int(col.get(qn('w:w')))
        except (TypeError, ValueError):
            return 0
    return total


def _scale_width(element, factor):
    """Kalikan atribut w:w sebuah gridCol/tcW dengan factor."""
    try:
        element.set(qn('w:w'), str(int(int(element.get(qn('w:w'))) * factor)))
    except (TypeError, ValueError):
        pass


def _widen_columns(table, grid, total_before):
    """
    Bagikan lebar kolom yang terhapus ke kolom yang tersisa, supaya tabel tetap
    memenuhi sel induknya dan tidak menyisakan ruang kosong di kanan.
    """
    total_after = _grid_width(grid)
    if not total_before or not total_after or total_after >= total_before:
        return

    factor = total_before / total_after
    for col in grid.findall(qn('w:gridCol')):
        _scale_width(col, factor)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.find(qn('w:tcPr'))
            tcW = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            if tcW is not None and tcW.get(qn('w:type')) == 'dxa':
                _scale_width(tcW, factor)


def _prune_unresolved_columns(table):
    """
    Hapus kolom yang header-nya masih berisi placeholder list yang tidak terisi.
    Dipakai oleh matriks seperti Korelasi CPMK-Sub-CPMK: template menyediakan
    kolom maksimum, context menentukan berapa yang benar-benar dipakai.
    Hanya untuk tabel tanpa merge, supaya struktur tidak rusak.
    """
    if not table.rows or _has_merged_cells(table):
        return
    doomed = [i for i, cell in enumerate(table.rows[0].cells)
              if _INDEXED_RE.search(cell.text)]
    if not doomed:
        return

    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    total_before = _grid_width(grid)

    for i in sorted(doomed, reverse=True):
        for tr in tbl.findall(qn('w:tr')):
            cells = tr.findall(qn('w:tc'))
            if i < len(cells):
                tr.remove(cells[i])
        if grid is not None:
            grid_cols = grid.findall(qn('w:gridCol'))
            if i < len(grid_cols):
                grid.remove(grid_cols[i])

    _widen_columns(table, grid, total_before)


# Nama array yang slot pertamanya (indeks 0) ternyata kosong. Hampir selalu
# berarti nama field di context tidak cocok dengan template, bukan section
# yang memang kosong. Dikumpulkan selama render, dilaporkan di akhir.
_EMPTY_ARRAYS = set()

# Array yang wajib ada isinya. Kalau kosong, dokumen hasil cacat berat --
# mis. 'detail' kosong berarti seluruh tabel mingguan hilang. Array di luar
# daftar ini (mis. 'korelasi') memang boleh tidak ada.
REQUIRED_ARRAYS = {'cpl_prodi', 'cpmk', 'sub_cpmk', 'detail', 'pustaka_utama'}


def _note_empty_arrays(placeholders):
    """Catat array yang bahkan elemen pertamanya tidak terisi."""
    for ph in placeholders:
        match = re.match(r'\{\s*([A-Za-z0-9_]+)\s*\[\s*0\s*\]', ph)
        if match:
            _EMPTY_ARRAYS.add(match.group(1))


def _is_unused_slot(before, after):
    """
    True bila teks tadinya berisi placeholder list dan tidak satu pun terisi.
    Artinya slot cadangan di template memang tidak dipakai oleh context ini.
    """
    if not before or not after or len(after) != len(before):
        return False
    return all(_INDEXED_RE.fullmatch(ph) for ph in after)


def _holds_empty_nested_table(row):
    """
    True bila baris ini memuat tabel bersarang yang tinggal baris header saja.
    Terjadi pada matriks korelasi ketika context belum memuat data terkait;
    seluruh blok (judul + matriks) dibuang, bukan disisakan kosong.
    """
    for cell in row.cells:
        for nested in cell.tables:
            if len(nested.rows) <= 1:
                return True
    return False


def _process_cell(cell, data, images):
    """
    Isi placeholder di satu sel, lalu buang paragraf slot yang tidak terpakai.
    Sel selalu menyisakan minimal satu paragraf, sesuai syarat format DOCX.
    """
    doomed = []
    for paragraph in cell.paragraphs:
        before = _PLACEHOLDER_RE.findall(paragraph.text)
        _process_paragraph(paragraph, data, images)
        if _is_unused_slot(before, _PLACEHOLDER_RE.findall(paragraph.text)):
            doomed.append(paragraph)

    if doomed and len(doomed) < len(cell.paragraphs):
        for paragraph in doomed:
            paragraph._p.getparent().remove(paragraph._p)

    for nested in cell.tables:
        _process_table(nested, data, images)


def _process_table(table, data: dict, images: dict):
    """
    Ganti placeholder di semua sel tabel, lalu buang baris cadangan yang tidak
    terpakai (template menyediakan lebih banyak baris daripada yang umumnya
    dibutuhkan) dan kolom yang tidak terpakai (mis. m9..m12 bila context hanya
    mengisi m1..m8).
    """
    for indeks, row in enumerate(list(table.rows)):
        before = _PLACEHOLDER_RE.findall(_row_text(row))
        for cell in row.cells:
            _process_cell(cell, data, images)
        after = _PLACEHOLDER_RE.findall(_row_text(row))
        # Baris pertama adalah header tabel; label kolomnya harus tetap ada
        # walau seluruh placeholder di dalamnya tidak terisi.
        if indeks == 0:
            continue
        if _is_unused_slot(before, after) or _holds_empty_nested_table(row):
            _note_empty_arrays(after)
            row._tr.getparent().remove(row._tr)

    _prune_unresolved_columns(table)


# ─── Public API ──────────────────────────────────────────────────────────────

def render(template_bytes: bytes, context: dict, images: dict = None) -> bytes:
    """
    Render template DOCX dengan context dict.

    Args:
        template_bytes : bytes dari file DOCX template
        context        : dict — keys matching placeholder names dalam template
        images         : dict — {'dosen_sign': bytes_png, 'mahasiswa_sign': bytes_png}

    Returns:
        bytes DOCX hasil render
    """
    if images is None:
        images = {}

    _EMPTY_ARRAYS.clear()
    doc = Document(io.BytesIO(template_bytes))

    for para in doc.paragraphs:
        _process_paragraph(para, context, images)

    for table in doc.tables:
        _process_table(table, context, images)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _process_paragraph(para, context, images)
        for para in section.footer.paragraphs:
            _process_paragraph(para, context, images)
        for table in section.header.tables:
            _process_table(table, context, images)
        for table in section.footer.tables:
            _process_table(table, context, images)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def get_missing_required_arrays() -> set:
    """
    Array wajib (lihat REQUIRED_ARRAYS) yang kosong sama sekali pada render
    terakhir. Panggil setelah render(); berguna untuk memperingatkan pengguna
    bila nama field context tidak cocok dengan template.
    """
    return _EMPTY_ARRAYS & REQUIRED_ARRAYS
